"""
Document ingestion: extract text from PDF/DOCX/TXT, chunk, embed, store.
"""
import io
import logging
from datetime import datetime
from pathlib import Path

from langchain_core.documents import Document

from app.core.config import get_settings
from app.services.vector_store import add_documents, get_vector_store
from app.utils.chunking import TextChunk, semantic_chunk

logger = logging.getLogger(__name__)

# Lazy imports for optional doc dependencies
_pypdf_available = False
_docx_available = False
try:
    from pypdf import PdfReader
    _pypdf_available = True
except ImportError:
    pass
try:
    from docx import Document as DocxDocument
    _docx_available = True
except ImportError:
    pass


def extract_text_from_pdf(content: bytes, filename: str) -> list[tuple[str, int | None]]:
    """Extract text per page from PDF. Returns [(text, page_number), ...]."""
    if not _pypdf_available:
        raise RuntimeError("PDF support requires pypdf. Install with: pip install pypdf")
    reader = PdfReader(io.BytesIO(content))
    result: list[tuple[str, int | None]] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        result.append((text, i + 1))
    return result


def extract_text_from_docx(content: bytes, filename: str) -> list[tuple[str, int | None]]:
    """Extract text from DOCX. Returns single block with page_number None."""
    if not _docx_available:
        raise RuntimeError("DOCX support requires python-docx. Install with: pip install python-docx")
    doc = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    text = "\n".join(parts)
    return [(text, None)]


def extract_text_from_txt(content: bytes, filename: str) -> list[tuple[str, int | None]]:
    """Extract text from TXT. Decode and return as single block."""
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1")
    return [(text, None)]


def extract_text(content: bytes, filename: str) -> list[tuple[str, int | None]]:
    """
    Dispatch by extension. Returns [(text, page_number), ...].
    Supported: .pdf, .docx, .txt
    """
    suf = Path(filename).suffix.lower()
    if suf == ".pdf":
        return extract_text_from_pdf(content, filename)
    if suf == ".docx":
        return extract_text_from_docx(content, filename)
    if suf == ".txt":
        return extract_text_from_txt(content, filename)
    raise ValueError(f"Unsupported file type: {suf}. Use .pdf, .docx, or .txt.")


def chunks_to_langchain_docs(chunks: list[TextChunk]) -> list[Document]:
    """Convert our TextChunk to LangChain Document with metadata."""
    return [
        Document(
            page_content=c.content,
            metadata={
                "document_name": c.document_name,
                "page_number": c.page_number,
                "upload_timestamp": c.upload_timestamp.isoformat(),
                "chunk_index": c.chunk_index,
            },
        )
        for c in chunks
    ]


async def ingest_file(content: bytes, filename: str) -> int:
    """
    Ingest one file: extract text, semantic chunk, add metadata, embed, store.
    Returns number of chunks created.
    """
    document_name = Path(filename).name
    upload_timestamp = datetime.utcnow()
    pages = extract_text(content, filename)
    all_chunks: list[TextChunk] = []
    for text, page_number in pages:
        if not text.strip():
            continue
        chunks = semantic_chunk(
            text,
            document_name=document_name,
            page_number=page_number,
            upload_timestamp=upload_timestamp,
        )
        all_chunks.extend(chunks)
    if not all_chunks:
        raise ValueError("No text could be extracted from the document.")
    docs = chunks_to_langchain_docs(all_chunks)
    add_documents(docs)
    return len(docs)
